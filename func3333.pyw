from MainWindow import Ui_MainWindow
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QPushButton, QLabel, QDialog, QLineEdit, QRadioButton, QGridLayout, QTableWidget, QTableWidgetItem, QFileDialog
from docx import Document
from docx.shared import Inches
from docx.text.run import Font, Run
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, os.path
from win32com import client
from PyQt5.QtGui import QPixmap
from docx.enum.style import WD_STYLE_TYPE
import json
import fitz


from PIL import Image
Header_1_Number = 1
Header_2_Number = 1
Header_3_Number = 1
listing_number = 1
Elements = []
ObjElements = []
picture_number = 1
table_number = 1
class Window(Ui_MainWindow):
    def __init__(self):
        super(Window, self).__init__()


    def UserUI(self):


        """===================================================================================================================="""

        self.ElementsGrid = QtWidgets.QVBoxLayout()
        self.TitleGroup = QtWidgets.QGroupBox()

        self.ElementsGrid.addStretch()
        self.ElementsGrid.setDirection(QtWidgets.QBoxLayout.BottomToTop)
        self.TitleGroup.setLayout(self.ElementsGrid)
        self.scrollArea_2.setWidget(self.TitleGroup)

        self.actionHeader_1.triggered.connect(lambda: self.Header_1())
        self.actionHeader_2.triggered.connect(lambda: self.Header_2())
        self.actionHeader_3.triggered.connect(lambda: self.Header_3())
        self.actionText.triggered.connect(lambda: self.Text())
        self.actionListing.triggered.connect(lambda: self.Listing())
        self.actionTable.triggered.connect(lambda: self.Table())
        self.actionList.triggered.connect(lambda: self.List())
        self.actionUsed_Info.triggered.connect(lambda: self.Usedinfo())
        self.actionPicture.triggered.connect(lambda: self.Picture())

        self.actionSave.triggered.connect(lambda: self.save_to())
        self.actionOpen.triggered.connect(lambda: self.open_from())

        self.Title_page()

    def save_to(self):
        path = []
        dialog = QFileDialog(None)
        dialog.setFileMode(QFileDialog.Directory)
        dialog.setViewMode(QFileDialog.List)

        if dialog.exec_():
            path = dialog.selectedFiles()

        if path:
            if os.path.exists(str(path[0]) + '/document.json'):
                file = open(str(path[0]) + '/document.json', 'w')
            else:
                file = open(str(path[0]) + '/document.json', 'x')
            json.dump(Elements, file)
            file.close()

    def open_from(self):
        path = []
        dialog = QFileDialog(None)
        dialog.setFileMode(QFileDialog.ExistingFile)
        dialog.setViewMode(QFileDialog.Detail)
        dialog.setNameFilter("Json files (*.json)")

        if dialog.exec_():
            path = dialog.selectedFiles()

        if path:
            file = open(str(path[0]), 'r')
            Elements.clear()
            Elements.extend(json.load(file))
            file.close()

            while self.ElementsGrid.count():
                w = self.ElementsGrid.takeAt(0)
                if w.widget():
                    w.widget().deleteLater()
            self.ElementsGrid.addStretch()

            for position in range(len(Elements)):
                if Elements[position][0] == "Header_1":
                    obj = Header_1_obj(position)
                    obj.HeaderText.setText(Elements[position][1])
                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] == "Header_2":
                    obj = Header_2_obj(position)
                    obj.HeaderText.setText(Elements[position][1])
                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] == "Header_3":
                    obj = Header_3_obj(position)
                    obj.HeaderText.setText(Elements[position][1])
                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] == "Text":
                    obj = Text_obj(position)
                    obj.textText.setPlainText(Elements[position][1])
                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] == "Listing":
                    obj = Listing_obj(position)
                    obj.listing_name.setText(Elements[position][1])
                    obj.listing_comment.setText(Elements[position][2])
                    obj.listingText.setPlainText(Elements[position][3])
                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] =="Table":
                    obj = Table_obj(position)

                    obj.table_name.setText(Elements[position][1])
                    obj.Table.setRowCount(Elements[position][2])
                    obj.Table.setColumnCount(Elements[position][3])

                    obj.height.setCurrentIndex(Elements[position][2] - 1)
                    obj.width.setCurrentIndex(Elements[position][3] - 1)

                    for i in range(Elements[position][2]):
                        for j in range(Elements[position][3]):
                            item = QTableWidgetItem(Elements[position][4][i][j])
                            obj.Table.setItem(i, j, item)

                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] == "Picture":
                    obj = Picture_obj(position)
                    self.ElementsGrid.insertWidget(1, obj.box)
                elif Elements[position][0] == "Title_page":
                    obj = Title_page_obj(position)

                    obj.workNameText.setText(Elements[position][1])
                    obj.themeText.setText(Elements[position][2])
                    obj.myNameText.setText(Elements[position][3])
                    obj.teacherText.setText(Elements[position][4])
                    obj.yearText.setText(Elements[position][5])

                    self.ElementsGrid.insertWidget(1, obj.box)


    def Header_1(self):
        Elements.append(["Header_1", ""])
        obj = Header_1_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)


    def Header_2(self):
        Elements.append(["Header_2", ""])
        obj = Header_2_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)


    def Header_3(self):
        Elements.append(["Header_3", ""])
        obj = Header_3_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)


    def Text(self):
        Elements.append(["Text", ""])
        obj = Text_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)


    def Title_page(self):
        Elements.append(["Title_page", "","","","",""])

        obj = Title_page_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)

        # self.ElementsGrid.insertWidget(1, Title_page_obj(len(Elements) - 1).box)



    def Listing(self):
        Elements.append(["Listing", "", "", ""])
        obj = Listing_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)


    def Usedinfo(self):
        Elements.append(["Usedinfo", ""])
        self.ElementsGrid.insertWidget(1, Usedinfo_obj(len(Elements) - 1).box)


    def Picture(self):
        Elements.append(["Picture", ""])
        self.ElementsGrid.insertWidget(1, Picture_obj(len(Elements) - 1).box)


    def Table(self):
        Elements.append(["Table", "", 0, 0, []])
        obj = Table_obj(len(Elements) - 1)
        ObjElements.append(obj)
        self.ElementsGrid.insertWidget(1, obj.box)


class Window4(QDialog):
    def __init__(self):
        super(Window4, self).__init__()


        self.setWindowTitle('Window4')
        self.setFixedSize(400,100)


        layoutqqq = QGridLayout(self)
        self.setLayout(layoutqqq)


        self.labqqq = QLabel(self)
        self.labqqq.setText("Сохранить как")
        layoutqqq.addWidget(self.labqqq, 0, 0)


        self.radioqqq1 = QRadioButton("docx")
        layoutqqq.addWidget(self.radioqqq1, 1, 0)


        self.radioqqq2 = QRadioButton("png")
        layoutqqq.addWidget(self.radioqqq2, 1, 1)


        self.btnqqq1 = QPushButton(self) 
        self.btnqqq1.setText("Отмена")
        self.btnqqq1.clicked.connect(self.reject)
        layoutqqq.addWidget(self.btnqqq1, 1, 2)


        self.btnqqq = QPushButton(self)
        self.btnqqq.setText("Продолжить")
        layoutqqq.addWidget(self.btnqqq,1,3)
        self.btnqqq.clicked.connect(self.chois)


    def chois(self):
        self.w7 = Window5()
        self.w7.exec()


class Window5(QDialog):


    def __init__(self):
        super(Window5, self).__init__()


        self.setWindowTitle('Window5')
        self.setFixedSize(400,100)


        layoutzzz = QGridLayout(self)
        self.setLayout(layoutzzz)


        self.labzzz = QLabel(self)
        self.labzzz.setText("Окно выбора места сохранения")
        layoutzzz.addWidget(self.labzzz, 0, 0)


        self.btnzzz = QPushButton(self)
        self.btnzzz.setText("Выбор")
        self.btnzzz.clicked.connect(self.showDialog)
        layoutzzz.addWidget(self.btnzzz, 1, 1)
        


        self.btnzzz1 = QPushButton(self) 
        self.btnzzz1.setText("Отмена")
        self.btnzzz1.clicked.connect(self.reject)
        layoutzzz.addWidget(self.btnzzz1, 1, 0)


class Title_page_obj():
    def __init__(self, position):
        self.position = position


        self.box = QtWidgets.QGroupBox('Титульный лист')
        self.box.setFixedHeight(450)
        self.layout = QtWidgets.QGridLayout()

        self.OkButton = QtWidgets.QPushButton('OK')

        self.lab = QLabel()
        self.lab.setText("Название работы")
        self.workNameText = QtWidgets.QLineEdit()

        self.lab1 = QLabel()
        self.lab1.setText("Тема работы")
        self.themeText = QtWidgets.QLineEdit()

        self.lab2 = QLabel()
        self.lab2.setText("Ваше ФИО")
        self.myNameText = QtWidgets.QLineEdit()

        self.lab5 = QLabel()
        self.lab5.setText("ФИО преподавателя")
        self.teacherText = QtWidgets.QLineEdit()

        self.lab8 = QLabel()
        self.lab8.setText("Год")
        self.yearText = QtWidgets.QLineEdit()


        self.layout.addWidget(self.lab, 0, 0)
        self.layout.addWidget(self.workNameText, 1, 0)
        self.layout.addWidget(self.lab1, 2, 0)
        self.layout.addWidget(self.themeText, 3, 0)
        self.layout.addWidget(self.lab2, 4, 0)
        self.layout.addWidget(self.myNameText, 5, 0)
        self.layout.addWidget(self.lab5, 6, 0)
        self.layout.addWidget(self.teacherText, 7, 0)
        self.layout.addWidget(self.lab8, 8, 0)
        self.layout.addWidget(self.yearText, 9, 0)



        self.layout.addWidget(self.OkButton, 10, 0)


        self.OkButton.clicked.connect(lambda: self.Title_page_ok_button_action(self.position, self.workNameText.text(), self.themeText.text(), self.myNameText.text(), self.teacherText.text(), self.yearText.text()))

        self.box.setLayout(self.layout)


        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")

    def Title_page_ok_button_action(self, position, value_name_work, value_theme_work, value_your_name, value_teacher_name, value_city_year):
        Elements[position][1] = value_name_work
        Elements[position][2] = value_theme_work
        Elements[position][3] = value_your_name
        Elements[position][4] = value_teacher_name
        Elements[position][5] = value_city_year
        add_to_docx_func()

    def add_to_docx(self, document):
        p=document.add_paragraph("Федеральное государственное автономное")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=document.add_paragraph("образовательное учреждение")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=document.add_paragraph("высшего образования")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=document.add_paragraph("«СИБИРСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ»")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph("")
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Институт космических и информационных технологий")
        style = document.styles['Normal']
        font = style.font
        font.name = 'Time New Romans'
        font.size = Pt(14)
        run.font.underline = True
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph1 = document.add_paragraph()
        run = paragraph1.add_run("институт")
        # style = document.styles['Normal']
        font = run.font
        font.name = 'Time New Romans'
        font.size = Pt(10)
        paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


        paragraph = document.add_paragraph()
        run = paragraph.add_run("Кафедра «Информатика»")
        style = document.styles['Normal']
        font = style.font
        font.name = 'Time New Romans'
        font.size = Pt(14)
        run.font.underline = True
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


        paragraph1 = document.add_paragraph()
        run = paragraph1.add_run("кафедра")
        # style = document.styles['Normal']
        font = run.font
        font.name = 'Time New Romans'
        font.size = Pt(10)
        paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER




        p=document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(self.workNameText.text())
        r.bold = True
        font = r.font
        font.name = 'Time New Romans'
        font.size = Pt(14)



        paragraph = document.add_paragraph("")
        paragraph = document.add_paragraph("")


        paragraph = document.add_paragraph()
        run = paragraph.add_run(self.themeText.text())
        style = document.styles['Normal']
        font = style.font
        font.name = 'Time New Romans'
        font.size = Pt(14)
        run.font.underline = True
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph1 = document.add_paragraph()
        run = paragraph1.add_run("тема")
        font = run.font
        font.name = 'Time New Romans'
        font.size = Pt(10)
        paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


        records = (
            ('Руководитель','','',self.myNameText.text()),
            ('','','подпись,дата','инициалы фамилия'),
            ('Студент','','', self.teacherText.text()),
            ('','номер зачетной книжки','подпись,дата','инициалы фамилия')
        )


        table = document.add_table(rows=4, cols=4)


        table.cell(0,0).width = 1097280
        table.cell(1,0).width = 1097280
        table.cell(2,0).width = 1097280
        table.cell(3,0).width = 1097280


        table.cell(0,1).width = 4846320 
        table.cell(1,1).width = 4846320 
        table.cell(2,1).width = 4846320 
        table.cell(3,1).width = 4846320 

        table.cell(0,2).width = 2846320 
        table.cell(1,2).width = 2846320 
        table.cell(2,2).width = 2846320 
        table.cell(3,2).width = 2846320


        table.cell(3,3).width = 2846320  

        for qty, i, desc, k in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = i
            row_cells[2].text = desc
            row_cells[3].text = k



        k = 0
        for row in table.rows:
            if(k % 2 == 1):
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(10)
            k += 1



        # p=document.add_paragraph(self.listingText.toPlainText())
        # p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p=document.add_paragraph("Красноярск " + self.yearText.text())
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # r=p.add_run("")
        # document.add_page_break()


class Header_1_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Заголовок 1 уровня')
        self.box.setFixedHeight(100)
        self.HeaderText = QtWidgets.QLineEdit()
        self.OkButton = QtWidgets.QPushButton('OK')
        self.OkButton.clicked.connect(lambda: self.Header_1_ok_button_action(self.position, self.HeaderText.text()))
        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))

        self.UpButton = QtWidgets.QPushButton('U')
        self.UpButton.setFixedSize(20, 20)
        self.UpButton.clicked.connect(lambda: up_action(ui, self))

        self.DownButton = QtWidgets.QPushButton('D')
        self.DownButton.setFixedSize(20, 20)
        self.DownButton.clicked.connect(lambda: down_action(ui, self))

        self.layout = QtWidgets.QGridLayout()
        self.layout.addWidget(self.UpButton, 0, 0)
        self.layout.addWidget(self.DownButton, 2, 0)
        self.layout.addWidget(self.HeaderText, 1, 1)
        self.layout.addWidget(self.OkButton, 1, 3)
        self.layout.addWidget(self.DeleteButton, 1, 4)

        self.box.setLayout(self.layout)

        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")

    def Header_1_ok_button_action(self, position, value):
        Elements[position][1] = value
        add_to_docx_func()


    def add_to_docx(self, document):
        global Header_1_Number
        global Header_2_Number
        global Header_3_Number
        Header_2_Number = 1
        Header_3_Number = 1


        run = document.add_heading(0).add_run(str(Header_1_Number) + " " + self.HeaderText.text())
        Header_1_Number += 1



        font = run.font
        font.name = "Times New Romans"

        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
 

class Header_2_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Заголовок 2 уровня')
        self.box.setFixedHeight(100)
        self.HeaderText = QtWidgets.QLineEdit()
        self.OkButton = QtWidgets.QPushButton('OK')
        self.OkButton.clicked.connect(lambda: self.Header_2_ok_button_action(self.position, self.HeaderText.text()))
        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))

        self.UpButton = QtWidgets.QPushButton('U')
        self.UpButton.setFixedSize(20, 20)
        self.UpButton.clicked.connect(lambda: up_action(ui, self))

        self.DownButton = QtWidgets.QPushButton('D')
        self.DownButton.setFixedSize(20, 20)
        self.DownButton.clicked.connect(lambda: down_action(ui, self))

        self.layout = QtWidgets.QGridLayout()
        self.layout.addWidget(self.UpButton, 0, 0)
        self.layout.addWidget(self.DownButton, 2, 0)
        self.layout.addWidget(self.HeaderText, 1, 1)
        self.layout.addWidget(self.OkButton, 1, 3)
        self.layout.addWidget(self.DeleteButton, 1, 4)

        self.box.setLayout(self.layout)

        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")
    def Header_2_ok_button_action(self, position, value):
        Elements[position][1] = value
        add_to_docx_func()


    def add_to_docx(self, document):
        global Header_1_Number
        global Header_2_Number
        global Header_3_Number

        Header_3_Number = 1


        if(Header_1_Number == 1):
        	Header_1_Number = 2



        run = document.add_heading(0).add_run(str(Header_1_Number-1)+"."+str(Header_2_Number) + " " +self.HeaderText.text())
        Header_2_Number += 1


        font = run.font
        font.name = "Times New Romans"

        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
 

class Header_3_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Заголовок 3 уровня')
        self.box.setFixedHeight(100)
        self.HeaderText = QtWidgets.QLineEdit()
        self.OkButton = QtWidgets.QPushButton('OK')
        self.OkButton.clicked.connect(lambda: self.Header_3_ok_button_action(self.position, self.HeaderText.text()))
        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))

        self.UpButton = QtWidgets.QPushButton('U')
        self.UpButton.setFixedSize(20, 20)
        self.UpButton.clicked.connect(lambda: up_action(ui, self))

        self.DownButton = QtWidgets.QPushButton('D')
        self.DownButton.setFixedSize(20, 20)
        self.DownButton.clicked.connect(lambda: down_action(ui, self))

        self.layout = QtWidgets.QGridLayout()
        self.layout.addWidget(self.UpButton, 0, 0)
        self.layout.addWidget(self.DownButton, 2, 0)
        self.layout.addWidget(self.HeaderText, 1, 1)
        self.layout.addWidget(self.OkButton, 1, 3)
        self.layout.addWidget(self.DeleteButton, 1, 4)

        self.box.setLayout(self.layout)


        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")
    def Header_3_ok_button_action(self, position, value):
        Elements[position][1] = value
        add_to_docx_func()

    def add_to_docx(self, document):
        global Header_1_Number
        global Header_2_Number
        global Header_3_Number

        if(Header_1_Number == 1):
        	Header_1_Number = 2
        if(Header_2_Number == 1):
        	Header_2_Number = 2

        run = document.add_heading(0).add_run(str(Header_1_Number-1) + "." + str(Header_2_Number-1) + "." + str(Header_3_Number) + " " + self.HeaderText.text())

        Header_3_Number += 1

        font = run.font
        font.name = "Times New Romans"

        font.color.rgb = RGBColor(0x00, 0x00, 0x00)


class Table_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Таблица')
        self.box.setFixedHeight(300)

        self.Table = QTableWidget()

        self.height = QtWidgets.QComboBox()
        self.height.setFixedWidth(40)
        self.height.addItems(['1', '2', '3', '4', '5', '6', '7', '8', '9'])
        self.width = QtWidgets.QComboBox()
        self.width.setFixedWidth(40)
        self.width.addItems(['1', '2', '3', '4', '5', '6', '7', '8', '9'])

        self.ResizeButton = QtWidgets.QPushButton('Resize')
        self.ResizeButton.clicked.connect(lambda: self.resize_func())
        self.OkButton = QtWidgets.QPushButton('OK')
        self.table0 = QtWidgets.QLabel("Введите название Таблицы")
        self.table_name = QtWidgets.QLineEdit()
        self.OkButton.clicked.connect(lambda: self.table_ok_button_action(self.position, self.Table.rowCount(), self.Table.columnCount(), self.table_name.text()))
        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))

        self.UpButton = QtWidgets.QPushButton('U')
        self.UpButton.setFixedSize(20, 20)
        self.UpButton.clicked.connect(lambda: up_action(ui, self))

        self.DownButton = QtWidgets.QPushButton('D')
        self.DownButton.setFixedSize(20, 20)
        self.DownButton.clicked.connect(lambda: down_action(ui, self))

        self.layout = QtWidgets.QGridLayout()
        self.layout.addWidget(self.UpButton, 2, 0)
        self.layout.addWidget(self.DownButton, 3, 0)
        self.layout.addWidget(self.table0, 0, 1)
        self.layout.addWidget(self.table_name, 1, 1, 1, 5)
        self.layout.addWidget(self.height, 2, 1)
        self.layout.addWidget(self.width, 2, 2)
        self.layout.addWidget(self.ResizeButton, 2, 3)
        self.layout.addWidget(self.OkButton, 2, 4)
        self.layout.addWidget(self.DeleteButton, 2, 5)
        self.layout.addWidget(self.Table, 3, 1, 1, 5)

        self.box.setLayout(self.layout)

        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")
    def resize_func(self):
        self.Table.setRowCount(int(self.height.currentText()))
        self.Table.setColumnCount(int(self.width.currentText()))

    def table_ok_button_action(self, position, rows, columns, name):
        Elements[position][4].clear()
        Elements[position][1] = name
        Elements[position][2] = rows
        Elements[position][3] = columns
        add_to_docx_func()
        for i in range(rows):
            Elements[position][4].append([])
            for j in range(columns):
                if self.Table.item(i, j) is not None:
                    Elements[position][4][i].append(self.Table.item(i, j).data(0))
                else:
                    Elements[position][4][i].append("")
        add_to_docx_func()


    def add_to_docx(self, document):
        try:
            global table_number
            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Romans'
            font.size = Pt(14)
            paragraph = document.add_paragraph("Таблица {0} – ".format(table_number) + self.table_name.text())
            # paragraph_format = paragraph.paragraph_format
            # paragraph_format.first_line_indent = Inches(0.5)
        

            table_number += 1
        except:
            pass
                            
        table = document.add_table(rows=self.Table.rowCount(), cols=self.Table.columnCount())
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(self.Table.rowCount()):
            for j in range(self.Table.columnCount()):
                if self.Table.item(i, j) is not None:
                    table.cell(i,j).text = self.Table.item(i, j).data(0)


class Text_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Обычный текст')
        self.box.setFixedHeight(120)

        self.textText = QtWidgets.QPlainTextEdit()
        self.OkButton = QtWidgets.QPushButton('OK')
        self.OkButton.clicked.connect(lambda: self.Text_ok_button_action(self.position, self.textText.toPlainText()))
        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))

        self.UpButton = QtWidgets.QPushButton('U')
        self.UpButton.setFixedSize(20, 20)
        self.UpButton.clicked.connect(lambda: up_action(ui, self))

        self.DownButton = QtWidgets.QPushButton('D')
        self.DownButton.setFixedSize(20, 20)
        self.DownButton.clicked.connect(lambda: down_action(ui, self))

        self.layout = QtWidgets.QGridLayout()
        self.layout.addWidget(self.UpButton, 0, 0)
        self.layout.addWidget(self.DownButton, 2, 0)
        self.layout.addWidget(self.textText, 0, 1, 3, 6)
        self.layout.addWidget(self.OkButton, 1, 7)
        self.layout.addWidget(self.DeleteButton, 1, 8)

        self.box.setLayout(self.layout)


        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")
    def Text_ok_button_action(self, position, value):
        Elements[position][1] = value
        add_to_docx_func()


    def add_to_docx(self, document):
        style = document.styles['Normal']
        font = style.font
        font.name = 'Time New Romans'
        font.size = Pt(14)
        paragraph = document.add_paragraph(self.textText.toPlainText())
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Inches(0.5)


class Usedinfo_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Список использованных источников')
        self.box.setFixedHeight(120)

        self.layout = QtWidgets.QGridLayout()


        self.listingText = QtWidgets.QPlainTextEdit()


        self.OkButton = QtWidgets.QPushButton('OK')


        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))


        self.layout.addWidget(self.listingText, 0, 0,1,0)
        self.layout.addWidget(self.OkButton, 1, 0)
        self.layout.addWidget(self.DeleteButton, 1, 1)



        self.OkButton.clicked.connect(lambda: self.Usedunfo_ok_button_action(self.position, self.listingText.toPlainText()))

        self.box.setLayout(self.layout)

        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")

    def Usedunfo_ok_button_action(self, position, value):
        Elements[position][1] = value


        p=document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run('СПИСОК ИСТОЧНИКОВ ИСТОЧНИКОВ')
        r.bold = True


        style = document.styles['Normal']
        font = style.font
        font.name = 'Time New Romans'
        font.size = Pt(14)
        paragraph = document.add_paragraph(value)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Inches(0.5)
        document.save('func.docx')


class Listing_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Листинг')
        self.box.setFixedHeight(240)


        self.OkButton = QtWidgets.QPushButton('OK')
        
        self.DeleteButton = QtWidgets.QPushButton('Delete')
        self.DeleteButton.clicked.connect(lambda: delete_button_action(ui, self))

        self.layout = QtWidgets.QGridLayout()


        self.text0 = QtWidgets.QLabel("Введите название листинга")
        self.listing_name = QtWidgets.QLineEdit()
        self.text = QtWidgets.QLabel("Введите комментарий")
        self.listing_comment = QtWidgets.QLineEdit()
        self.text1 = QtWidgets.QLabel("Введите листинг")
        self.listingText = QtWidgets.QPlainTextEdit()

        self.UpButton = QtWidgets.QPushButton('U')
        self.UpButton.setFixedSize(20, 20)
        self.UpButton.clicked.connect(lambda: up_action(ui, self))

        self.DownButton = QtWidgets.QPushButton('D')
        self.DownButton.setFixedSize(20, 20)
        self.DownButton.clicked.connect(lambda: down_action(ui, self))

        self.layout.addWidget(self.UpButton, 2, 0)
        self.layout.addWidget(self.DownButton, 4, 0)
        self.layout.addWidget(self.text0, 0, 1)
        self.layout.addWidget(self.listing_name, 1, 1)
        self.layout.addWidget(self.text, 2, 1)
        self.layout.addWidget(self.listing_comment, 3, 1)
        self.layout.addWidget(self.text1, 4, 1)
        self.layout.addWidget(self.listingText, 5, 1)
        self.layout.addWidget(self.OkButton, 6, 1)
        self.layout.addWidget(self.DeleteButton, 6, 2)


        self.OkButton.clicked.connect(lambda: self.Listing_ok_button_action(self.position, self.listing_name.text(), self.listing_comment.text(), self.listingText.toPlainText()))

        self.box.setLayout(self.layout)

        
        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")


    def Listing_ok_button_action(self, position, value_name, value_comment, value_text):
        Elements[position][1] = value_name
        Elements[position][2] = value_comment
        Elements[position][3] = value_text
        add_to_docx_func()


    def add_to_docx(self, document):
        global listing_number
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Листинг {0} - ".format(listing_number) + self.listing_name.text())
        listing_number += 1
        style = document.styles['Normal']
        font = style.font
        font.name = 'Time New Romans'
        font.size = Pt(14)


        paragraph1 = document.add_paragraph()
        run = paragraph1.add_run("#"+self.listing_comment.text())
        # style = document.styles['Normal']
        font = run.font
        font.name = 'Courier New'
        font.size = Pt(10)


        paragraph1 = document.add_paragraph()
        run = paragraph1.add_run(self.listingText.toPlainText())
        # style = document.styles['Normal']
        font = run.font
        font.name = 'Courier New'
        font.size = Pt(10)

        # paragraph = document.add_paragraph(self.listingText.toPlainText())
        # paragraph_format = paragraph.paragraph_format
        # paragraph_format.first_line_indent = Inches(0.5)


class Picture_obj():
    def __init__(self, position):
        self.position = position

        self.box = QtWidgets.QGroupBox('Рисунок')
        self.box.setFixedHeight(300)

        self.ChooseButton = QtWidgets.QPushButton('Choose picture')
        self.draw = QtWidgets.QLabel("Название рисунка")
        self.draw_choice = QtWidgets.QLineEdit()
        self.ChooseButton.clicked.connect(lambda: self.choose_picture_action(self.position, self.draw_choice.text()))
        self.Picture = QPixmap()

        self.label = QLabel()

        self.layout = QtWidgets.QGridLayout()
        self.layout.addWidget(self.draw, 0, 0)
        self.layout.addWidget(self.draw_choice, 1,0)
        self.layout.addWidget(self.ChooseButton, 2, 4, 5, 5)
        self.layout.addWidget(self.label, 3, 0, 1, 4)
        self.box.setLayout(self.layout)


        self.box.setStyleSheet("font-weight:bold;\n"
"font-family:lato;\n")

    def choose_picture_action(self, position, value):
        path = []
        dialog = QFileDialog(None)
        dialog.setFileMode(QFileDialog.ExistingFile)
        dialog.setNameFilter("Images (*.png *.jpg *.jpeg)")
        dialog.setViewMode(QFileDialog.Detail)

        if dialog.exec_():
            path = dialog.selectedFiles()

        if path:
            self.Picture.load(path[0])
            self.Picture = self.Picture.scaled(350, 350, QtCore.Qt.KeepAspectRatio)
            self.label.setPixmap(self.Picture)


        try:
            global picture_number
            Elements[position][1] = value
            paragraph = document.add_picture(path[0], width=Inches(5.25), height=Inches(4.7))

            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


            style = document.styles['Normal']
            font = style.font
            font.name = 'Time New Romans'
            font.size = Pt(14)
            paragraph = document.add_paragraph("Рисунок {0} – ".format(picture_number) + value)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.first_line_indent = Inches(0.5)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            picture_number += 1

            document.save('func.docx')
        except:
            pass


def test():
    assert (listing_number>0 and picture_number>0) == True


def add_to_docx_func():
    global Header_1_Number
    Header_1_Number = 1
    global Header_2_Number
    Header_2_Number = 1
    global Header_3_Number
    Header_3_Number = 1
    global listing_number
    listing_number = 1
    global picture_number
    picture_number = 1
    global table_number
    table_number = 1

    document = Document()
    for obj in ObjElements:
        obj.add_to_docx(document)
    document.save('func.docx')
    preview_img()
    preview(ui)


def preview_img():
    file_type = 'docx'
    try:
        for files in os.listdir("."):
            if files.endswith(".docx"):
                out_name = files.replace(file_type, r"pdf")
                in_file = os.path.abspath(files)
                out_file = os.path.abspath(out_name)
                word = client.Dispatch("Word.Application")
                doc = word.Documents.Open(in_file)
                doc.SaveAs(out_file, FileFormat=17)
                doc.Close()
                word.Quit()
    except:
        print("error")



    count = 0
    doc = fitz.open("func.pdf")
    for page in doc:
        pix = page.getPixmap()
        pix.writeImage("func-%i.png" % page.number)
        count += 1



    if(count>0):
        if(count!=1):
            img = Image.new('RGB', (595, 842*count))
            for i in range(count):
                a = Image.open("{0}.png".format("func-"+str(i)))
                img.paste(a, (0,(842*i)))
        else:
            img = Image.new('RGB', (595, 842*count))
            img0 = Image.open("func-0.png")
            img.paste(img0, (0,0))

        img.save("out.png")
    else:
        pass


def preview(window):
    grid = QtWidgets.QGridLayout()
    group = QtWidgets.QGroupBox()
    group.setLayout(grid)
    window.scrollArea.setWidget(group)

    count = 0
    for files in os.listdir("."):
        if files.endswith("func-%i.png" % count):
            count += 1

    picture = QLabel()
    picture.setPixmap(QPixmap("out.png"))
    picture.setFixedSize(595, count*842)
    picture.setScaledContents(True)

    grid.addWidget(picture)


    for i in range(count):
        os.remove("func-%i.png" % i)
    os.remove("out.png")


def delete_button_action(window, obj):
    for i in range(1, len(Elements) - obj.position):
        ObjElements[obj.position + i].position -= 1

    del Elements[obj.position]
    del ObjElements[obj.position]
    window.ElementsGrid.removeWidget(obj.box)
    obj.box.deleteLater()
    obj = None
    add_to_docx_func()


def up_action(window, obj):
    if obj.position > 1:
        window.ElementsGrid.insertWidget(len(Elements) - obj.position + 1, obj.box)

        ObjElements[obj.position - 1], ObjElements[obj.position] = ObjElements[obj.position], ObjElements[obj.position - 1]
        Elements[obj.position - 1], Elements[obj.position] = Elements[obj.position], Elements[obj.position - 1]

        ObjElements[obj.position].position += 1
        obj.position -= 1


def down_action(window, obj):
    if obj.position < len(Elements) - 1:
        window.ElementsGrid.insertWidget(len(Elements) - obj.position - 1, obj.box)

        ObjElements[obj.position + 1], ObjElements[obj.position] = ObjElements[obj.position], ObjElements[obj.position + 1]
        Elements[obj.position + 1], Elements[obj.position] = Elements[obj.position], Elements[obj.position + 1]

        ObjElements[obj.position].position -= 1
        obj.position += 1




    else:
        pass




if __name__ == "__main__":
    import sys
    test()
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Window()
    ui.setupUi(MainWindow)
    ui.UserUI()
    MainWindow.show()
    sys.exit(app.exec_())
